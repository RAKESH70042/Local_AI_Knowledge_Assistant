# backend/main.py  — Full RAG backend
import os
from fastapi.responses import StreamingResponse
from core.redactor import redact
from fastapi import Depends, FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from pydantic import BaseModel
from typing import Optional, List
import time, json, shutil
from pathlib import Path
from datetime import datetime
from fastapi.security import APIKeyHeader
from fastapi import Security
from dotenv import load_dotenv
load_dotenv()  
# MUST be before any os.getenv() call
# FIX 1: warn loudly on startup if token is still the default
API_TOKEN = os.getenv("API_TOKEN", "")
if not API_TOKEN:
    raise RuntimeError(
        "API_TOKEN is not set in your .env file.\n"
        "Add:  API_TOKEN=my-secret-key-123"
    )

api_key_header = APIKeyHeader(name="X-API-Token", auto_error=False)

# FIX 2: verify_token must return the key on success so FastAPI Depends works correctly
async def verify_token(key: str = Security(api_key_header)):
    if key != API_TOKEN:
        raise HTTPException(status_code=403, detail="Invalid or missing API token")
    return key  # <-- was missing, caused unpredictable auth behaviour

# ── Try to import your existing core modules ──────────────
try:
    from core.pipeline import RAGPipeline
    _has_pipeline = True
except Exception as e:
    print(f"[WARN] Could not import RAGPipeline: {e}")
    _has_pipeline = False

try:
    from core.database import ChatDatabase
    _has_db = True
except Exception as e:
    print(f"[WARN] Could not import ChatDatabase: {e}")
    _has_db = False

# ── App ───────────────────────────────────────────────────
app = FastAPI(title="AI Knowledge Assistant", version="2.0.0")

# FIX 3: restrict CORS to localhost only — "*" allows any website to call your API
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://127.0.0.1:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

BASE_DIR  = Path(__file__).parent.parent
DOCS_DIR  = BASE_DIR / "data" / "docs"
INDEX_DIR = BASE_DIR / "data" / "index"
LOG_FILE  = BASE_DIR / "data" / "query_log.jsonl"
DOCS_DIR.mkdir(parents=True, exist_ok=True)
INDEX_DIR.mkdir(parents=True, exist_ok=True)
LOG_FILE.parent.mkdir(parents=True, exist_ok=True)

pipeline = RAGPipeline(docs_dir=DOCS_DIR, index_dir=INDEX_DIR) if _has_pipeline else None
db       = ChatDatabase(BASE_DIR / "data" / "chat_history.db") if _has_db else None

# ── Models ────────────────────────────────────────────────
class AskRequest(BaseModel):
    question: str
    top_k: Optional[int] = 5

class ChatRequest(BaseModel):
    message: str
    session_id: Optional[str] = "default"
    history: Optional[List[dict]] = []
    context: Optional[str] = None
    filename: Optional[str] = None

# ── Health ────────────────────────────────────────────────
@app.get("/")
@app.get("/health")
async def health():
    stats = pipeline.get_stats() if pipeline else {}
    return {
        "status": "ok",
        "message": "AI Knowledge Assistant is running!",
        "indexed": stats.get("indexed", False),
        "total_chunks": stats.get("total_chunks", 0),
        "llm": f"ollama/{os.getenv('OLLAMA_MODEL', 'qwen2.5:1.5b')}",
        "timestamp": datetime.utcnow().isoformat(),
    }

# ── /ask  (core RAG endpoint) ─────────────────────────────
@app.post("/ask")
async def ask(req: AskRequest, _=Depends(verify_token)):
    if not req.question.strip():
        raise HTTPException(400, "Question cannot be empty")
    if not pipeline:
        raise HTTPException(503, "Pipeline not initialised — check backend logs")

    t0 = time.perf_counter()
    result = await pipeline.ask(req.question, top_k=req.top_k or 5)
    latency_ms = round((time.perf_counter() - t0) * 1000, 1)

    result.setdefault("metrics", {})
    result["metrics"]["latency_ms"] = latency_ms
    result["metrics"]["chunk_ids"]  = result.get("sources", [])

    result["answer"] = redact(result["answer"])
    _log(req.question, result, latency_ms)
    return result

# ── /chat  (frontend-compatible wrapper around /ask) ──────
@app.post("/chat")
async def chat(req: ChatRequest, _=Depends(verify_token)):
    result = await ask(AskRequest(question=req.message))

    if db:
        try:
            db.save_message(req.session_id, "user",      req.message)
            db.save_message(req.session_id, "assistant", result["answer"])
        except Exception:
            pass

    return {
        "response": result["answer"],
        "answer":   result["answer"],
        "sources":  result.get("sources", []),
        "metrics":  result.get("metrics", {}),
    }

# ── /stream ───────────────────────────────────────────────
@app.post("/stream")
async def stream(req: ChatRequest, _=Depends(verify_token)):
    if not pipeline:
        raise HTTPException(503, "Pipeline not available")

    from core.embeddings  import embed_query
    from core.vectorstore import search_chunks
    from core.llm         import ask_llm_stream
    import asyncio

    async def generate():
        try:
            chunks = []

            #  Use uploaded file context if provided
            if req.context:
                chunks = [{
                    "text":   req.context[:6000],  # trim if huge
                    "source": req.filename or "uploaded file",
                    "score":  1.0,
                }]
            else:
                # Fall back to vector search
                query_embedding = embed_query(req.message)
                chunks = search_chunks(query_embedding, top_k=5)

            if not chunks:
                yield " No relevant documents found. Please index some documents first."
                return

            for token in ask_llm_stream(req.message, chunks):
                yield token
                await asyncio.sleep(0)

        except Exception as e:
            yield f" Error: {str(e)}"

    return StreamingResponse(generate(), media_type="text/plain")

# ── /ingest ───────────────────────────────────────────────
@app.post("/ingest")
async def ingest(force: bool = False):
    if not pipeline:
        raise HTTPException(503, "Pipeline not available")
    t0 = time.perf_counter()
    result = pipeline.ingest(force=force)
    result["elapsed_sec"] = round(time.perf_counter() - t0, 2)
    return result

# ── /upload ───────────────────────────────────────────────
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".gif"}

@app.post("/upload")
async def upload(file: UploadFile = File(...)):
    allowed = {".pdf", ".docx", ".doc", ".txt", ".md", ".csv"} | IMAGE_EXTS
    ext     = Path(file.filename).suffix.lower()

    if ext not in allowed:
        raise HTTPException(400, f"Type '{ext}' not supported.")

    stamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    dest  = DOCS_DIR / f"{stamp}_{file.filename}"

    with open(dest, "wb") as f:
        shutil.copyfileobj(file.file, f)

    # ── Extract text ──────────────────────────────────────
    extracted_text = ""
    if ext in IMAGE_EXTS:
        extracted_text = _ocr_image(str(dest))
    else:
        from core.parser import parse_file
        try:
            extracted_text = parse_file(str(dest))
        except Exception:
            extracted_text = ""

    # Images aren't indexed into the vector store (no meaningful chunks)
    ingest_result = {}
    if ext not in IMAGE_EXTS:
        ingest_result = pipeline.ingest_file(str(dest), dest.name) if pipeline else {}

    return {
        "status":         "uploaded",
        "filename":       dest.name,
        "original_name":  file.filename,
        "size_kb":        round(dest.stat().st_size / 1024, 1),
        "extracted_text": extracted_text[:6000],
        "reindex":        ingest_result,
    }


## ── OCR helper (Ollama vision — moondream) ────────────────
def _ocr_image(path: str) -> str:
    """Extract text from an image using Ollama vision model."""
    import base64, httpx

    OLLAMA_URL   = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
    VISION_MODEL = os.getenv("OLLAMA_VISION_MODEL", "moondream")

    try:
        with open(path, "rb") as f:
            image_b64 = base64.b64encode(f.read()).decode("utf-8")

        response = httpx.post(
            f"{OLLAMA_URL}/api/generate",
            json={
                "model":  VISION_MODEL,
                "prompt": "Please read and transcribe every word of text visible in this image. Include the person's name, contact details, profile, education, skills, and certifications exactly as written.",
                "images": [image_b64],
                "stream": False,
            },
            timeout=60.0,
        )
        response.raise_for_status()
        text = response.json().get("response", "").strip()
        return text if text else "[No text detected in image]"

    except Exception as e:
        return f"[Vision OCR failed: {e}]"

# ── /documents ────────────────────────────────────────────
@app.get("/documents")
async def list_documents():
    allowed = {".pdf", ".docx", ".doc", ".txt", ".md", ".csv", ".odt"}
    docs = [
        {
            "name":    p.name,
            "size_kb": round(p.stat().st_size / 1024, 1),
            "type":    p.suffix.lstrip(".").upper(),
        }
        for p in sorted(DOCS_DIR.rglob("*"))
        if p.is_file() and p.suffix.lower() in allowed
    ]
    return {
        "documents":   docs,
        "total":       len(docs),
        "index_stats": pipeline.get_stats() if pipeline else {},
    }

# ── /documents/{filename} ─────────────────────────────────
@app.delete("/documents/{filename}")
async def delete_doc(filename: str):
    target = DOCS_DIR / filename
    if not target.exists():
        raise HTTPException(404, "File not found")
    target.unlink()
    # Remove only this file's chunks from ChromaDB — don't re-index everything
    from core.vectorstore import delete_source
    delete_source(filename)
    return {"deleted": filename, "status": "chunks removed from index"}

# ── /metrics ──────────────────────────────────────────────
@app.get("/metrics")
async def get_metrics():
    entries = _read_log()
    if not entries:
        return {"total_queries": 0}
    lats = [e["latency_ms"] for e in entries if "latency_ms" in e]
    return {
        "total_queries":  len(entries),
        "avg_latency_ms": round(sum(lats) / len(lats), 1) if lats else 0,
        "recent_queries": [
            {
                "timestamp":  e.get("timestamp"),
                "query":      e.get("query", "")[:80],
                "latency_ms": e.get("latency_ms"),
                "sources":    e.get("sources", []),
            }
            for e in entries[-10:][::-1]
        ],
    }

# ── /export/docx ──────────────────────────────────────────
@app.post("/export/docx")
async def export_docx(req: AskRequest):
    result = await pipeline.ask(req.question) if pipeline else {"answer": "", "sources": []}
    data   = _make_docx(req.question, result["answer"], result.get("sources", []))
    fname  = f"answer_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.docx"
    return Response(
        content=data,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )

# ── /export/xlsx ──────────────────────────────────────────
@app.get("/export/xlsx")
async def export_xlsx():
    data  = _make_xlsx(_read_log())
    fname = f"query_log_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.xlsx"
    return Response(
        content=data,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )

# ── /history ──────────────────────────────────────────────
@app.get("/history")
async def get_history():
    from core.database import get_all_chats
    return {"history": get_all_chats()}

@app.delete("/history")
async def clear_history():
    from core.database import clear_all_chats
    clear_all_chats()
    return {"status": "cleared"}


# ── Internal helpers ──────────────────────────────────────

def _log(query, result, latency_ms):
    entry = {
        "timestamp":  datetime.utcnow().isoformat() + "Z",
        "query":      query,
        "latency_ms": latency_ms,
        "sources":    result.get("sources", []),
        "chunks":     result.get("metrics", {}).get("chunk_ids", []),
    }
    with open(LOG_FILE, "a") as f:
        f.write(json.dumps(entry) + "\n")


def _read_log():
    if not LOG_FILE.exists():
        return []
    out = []
    for line in LOG_FILE.read_text().splitlines():
        try:
            out.append(json.loads(line))
        except Exception:
            pass
    return out


def _make_docx(question: str, answer: str, sources: list) -> bytes:
    import io
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("AI Knowledge Assistant — Answer", 1)
    doc.add_paragraph(f"Exported: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")

    doc.add_heading("Question", 2)
    p = doc.add_paragraph(question)
    p.runs[0].bold = True

    doc.add_heading("Answer", 2)
    for line in answer.split("\n"):
        doc.add_paragraph(line)

    if sources:
        doc.add_heading("Sources / Citations", 2)
        for s in sources:
            # FIX 4: sources is a list of plain strings (filenames), not dicts.
            # Old code called s.get("source") which crashes with AttributeError.
            if isinstance(s, dict):
                # handle dict form just in case pipeline ever returns dicts
                label   = s.get("source", "Unknown")
                excerpt = s.get("text", "")
            else:
                # normal case: s is a filename string like "synth_fennex_overview.txt"
                label   = s
                excerpt = ""

            doc.add_paragraph(label, style="List Bullet")

            if excerpt:
                ex = doc.add_paragraph(f'"{excerpt[:200]}…"')
                ex.runs[0].italic    = True
                ex.runs[0].font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_xlsx(entries: list) -> bytes:
    import io, xlsxwriter

    buf = io.BytesIO()
    wb  = xlsxwriter.Workbook(buf)
    ws  = wb.add_worksheet("Query Log")
    hf  = wb.add_format({"bold": True, "bg_color": "#1A56DB", "font_color": "white", "border": 1})
    cf  = wb.add_format({"text_wrap": True, "valign": "top", "border": 1})

    headers = ["#", "Query", "Sources", "Latency (ms)", "Timestamp"]
    widths  = [4, 55, 35, 13, 22]
    for c, (h, w) in enumerate(zip(headers, widths)):
        ws.write(0, c, h, hf)
        ws.set_column(c, c, w)

    for r, e in enumerate(entries, 1):
        src = ", ".join(e.get("sources", [])) if isinstance(e.get("sources"), list) else ""
        ws.write(r, 0, r,                      cf)
        ws.write(r, 1, e.get("query",      ""), cf)
        ws.write(r, 2, src,                     cf)
        ws.write(r, 3, e.get("latency_ms", 0),  cf)
        ws.write(r, 4, e.get("timestamp",  ""), cf)
        ws.set_row(r, 36)

    wb.close()
    buf.seek(0)
    return buf.getvalue()